"""
Invoice Generator Service - Handles multi-page invoices with pagination
"""
from pathlib import Path
from io import BytesIO
from datetime import datetime
from copy import deepcopy
import re
import subprocess
from decimal import Decimal
import logging
import os
import json

from docxtpl import DocxTemplate
from docx import Document
from docx.oxml.shared import OxmlElement
from docx.shared import Pt
from core.database import supabase
from fastapi import HTTPException
from services.utils import convert_docx_to_pdf

logger = logging.getLogger(__name__)

# Create output directory for invoices
OUTPUT_DIR = Path(__file__).parent.parent / "output" / "invoices"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
logger.info(f"Invoice output directory: {OUTPUT_DIR}")


# Bank details (constant)
BANK_DETAILS = {
    "ac_name": "ASK INTERNATIONAL",
    "ac_no": "7041-0212000820",
    "bank_name": "TRUST BANK LIMITED",
    "branch_name": "KAFRUL BRANCH",
    "routing_number": "240262387"
}

ITEMS_PER_PAGE = 30


def money(x) -> str:
    """Format number with thousand separators."""
    try:
        f = float(x)
        return f"{int(f):,}" if f.is_integer() else f"{f:,.2f}"
    except Exception:
        return str(x)


def parse_date_ddmmyyyy(s: str) -> str:
    """Normalize date to DD/MM/YYYY format."""
    for fmt in ("%m/%d/%Y", "%m/%d/%y", "%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            dt = datetime.strptime(s, fmt)
            return dt.strftime("%d/%m/%Y")
        except Exception:
            continue
    try:
        parts = s.replace("-", "/").split("/")
        if len(parts) == 3:
            m, d, y = parts
            dt = datetime.strptime(f"{int(m)}/{int(d)}/{int(y)}", "%m/%d/%Y")
            return dt.strftime("%d/%m/%Y")
    except Exception:
        pass
    return s


def number_to_words(num):
    """Convert number to words (Bangladeshi Taka)."""
    try:
        num = float(num)
        taka = int(num)
        paisa = int(round((num - taka) * 100))

        ones = ['', 'One', 'Two', 'Three', 'Four', 'Five', 'Six', 'Seven', 'Eight', 'Nine']
        teens = ['Ten', 'Eleven', 'Twelve', 'Thirteen', 'Fourteen', 'Fifteen',
                 'Sixteen', 'Seventeen', 'Eighteen', 'Nineteen']
        tens = ['', '', 'Twenty', 'Thirty', 'Forty', 'Fifty', 'Sixty', 'Seventy', 'Eighty', 'Ninety']

        def convert_below_thousand(n):
            if n == 0:
                return ''
            if n < 10:
                return ones[n]
            if n < 20:
                return teens[n - 10]
            if n < 100:
                return tens[n // 10] + (' ' + ones[n % 10] if n % 10 else '')
            return ones[n // 100] + ' Hundred' + (' ' + convert_below_thousand(n % 100) if n % 100 else '')

        if taka == 0:
            result = 'Zero Taka'
        else:
            crore = taka // 10000000
            lakh = (taka % 10000000) // 100000
            thousand = (taka % 100000) // 1000
            hundred = taka % 1000

            parts = []
            if crore:   parts.append(convert_below_thousand(crore) + ' Crore')
            if lakh:    parts.append(convert_below_thousand(lakh) + ' Lakh')
            if thousand: parts.append(convert_below_thousand(thousand) + ' Thousand')
            if hundred: parts.append(convert_below_thousand(hundred))
            result = (' '.join(parts) + ' Taka').strip()

        if paisa:
            result += ' and ' + convert_below_thousand(paisa) + ' Paisa'

        return result.strip() + ' Only'
    except Exception:
        return "Amount in Words"


def set_cell_font(cell, font_name='Arial', font_size=8):
    """Set font name and size for all runs in a cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)


def add_table_row_from_template(table, row_data: list, template_row):
    """Clone template row and replace text, preserving formatting."""
    try:
        w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        W_T = f'{{{w_ns}}}t'
        W_P = f'{{{w_ns}}}p'
        W_R = f'{{{w_ns}}}r'
        W_TC = f'{{{w_ns}}}tc'

        new_tr = deepcopy(template_row._element)
        table._tbl.append(new_tr)

        cells = new_tr.findall(f'.//{W_TC}')
        for idx, cell_text in enumerate(row_data):
            if idx >= len(cells):
                break
            cell = cells[idx]
            t_elems = list(cell.iter(W_T))

            if t_elems:
                t_elems[0].text = str(cell_text)
                for t in t_elems[1:]:
                    t.text = ''
            else:
                first_p = cell.find(f'.//{W_P}')
                if first_p is None:
                    first_p = OxmlElement('w:p')
                    cell.append(first_p)
                first_r = first_p.find(f'.//{W_R}')
                if first_r is None:
                    first_r = OxmlElement('w:r')
                    first_p.append(first_r)
                t = OxmlElement('w:t')
                t.text = str(cell_text)
                first_r.append(t)

        new_row = table.rows[-1]
        for cell in new_row.cells:
            set_cell_font(cell, 'Arial', 8)

        return True
    except Exception as e:
        print(f"Error adding row: {e}")
        return False


def find_items_table(doc: Document):
    """Locate the items table by matching header cells."""
    wanted = ["Sl #", "Product Description", "Pkt Size", "Qty", "Unit Price", "Total Price"]
    for tbl in doc.tables:
        try:
            hdr = [cell.text.strip() for cell in tbl.rows[0].cells]
            if len(hdr) >= 6 and all(h in hdr for h in wanted):
                return tbl
        except Exception:
            pass
    return None


def replace_placeholders_everywhere(doc: Document, context: dict):
    """Replace placeholders in all text nodes."""
    try:
        repl = {f"{{{{ {k} }}}}": str(v) for k, v in context.items()}
        repl.update({f"{{{{{k}}}}}": str(v) for k, v in context.items()})

        W_T = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
        root = doc.part.element

        for t in root.iter(W_T):
            if t.text:
                orig = t.text
                new = orig
                for k, v in repl.items():
                    if k in new:
                        logger.info(f"Replacing placeholder {k} with {v}")
                        new = new.replace(k, v)
                if new != orig:
                    t.text = new
        return True
    except Exception as e:
        logger.error(f"Error replacing placeholders: {e}")
        return False




class InvoiceGeneratorService:
    """Generate multi-page invoices with pagination."""
    
    @staticmethod
    def generate_invoice_for_po(po_id: int, template_path: Path, output_dir: Path = None) -> tuple:
        """
        Generate invoice for a purchase order.
        Returns: (docx_path, pdf_path)
        """
        logger.info(f"Starting invoice generation for PO ID: {po_id}")
        
        # Use persistent output directory if not specified
        if output_dir is None:
            output_dir = OUTPUT_DIR
        
        output_dir.mkdir(parents=True, exist_ok=True)
        logger.debug(f"Output directory: {output_dir}")
        
        # Get PO details
        logger.info(f"Fetching purchase order details for PO ID: {po_id}")
        po_res = supabase.table("purchase_orders").select("*").eq("po_id", po_id).execute()
        if not po_res.data:
            logger.error(f"Purchase order not found for PO ID: {po_id}")
            raise HTTPException(status_code=404, detail="Purchase order not found")
        
        po = po_res.data[0]
        logger.info(f"PO retrieved: {po.get('po_number')} - Status: {po.get('status')}")
        
        # Get dealer details
        logger.info(f"Fetching dealer details for dealer ID: {po['dealer_id']}")
        dealer_res = supabase.table("dealers").select("*").eq("dealer_id", str(po["dealer_id"])).execute()
        if not dealer_res.data:
            logger.error(f"Dealer not found for dealer ID: {po['dealer_id']}")
            raise HTTPException(status_code=404, detail="Dealer not found")
        
        dealer = dealer_res.data[0]
        logger.info(f"Dealer retrieved: {dealer.get('company_name')} (Code: {dealer.get('customer_code')})")
        
        # Get PO items
        logger.info(f"Fetching items for PO ID: {po_id}")
        items_res = supabase.table("purchase_order_items").select("*").eq("po_id", po_id).execute()
        items = items_res.data or []
        logger.info(f"Retrieved {len(items)} items for PO")
        
        # Fetch product details for each item
        logger.info(f"Fetching product details for items")
        if items:
            product_ids = list({it["product_id"] for it in items if it.get("product_id")})
            if product_ids:
                pres = supabase.table("products").select("*").in_("product_id", product_ids).execute()
                product_map = {p["product_id"]: p for p in (pres.data or [])}
            else:
                product_map = {}
            
            # Inject product details into items
            for it in items:
                it["product"] = product_map.get(it.get("product_id"))
        
        # Prepare items with totals
        items_with_totals = []
        subtotal = 0.0
        
        for idx, item in enumerate(items, 1):
            qty = float(item.get("quantity", 0))
            
            # Get product details
            product = item.get("product", {})
            product_name = product.get("name", "") if product else ""
            pack_size = product.get("pack_size", "") if product else ""
            
            # Calculate unit_price from product's trade_price_incl_vat
            unit_price = float(product.get("TP", 0)) if product else 0.0
            total = qty * unit_price
            subtotal += total
            
            items_with_totals.append({
                "sl": idx,
                "desc": product_name,
                "pkt_size": pack_size,
                "qty": qty,
                "unit_price": unit_price,
                "total": total,
            })
            logger.debug(f"Item {idx}: {product_name} - Qty: {qty}, Unit Price: {unit_price}, Total: {total}")
        
        logger.info(f"Subtotal calculated: {subtotal}")
        
        # Fetch VAT and commission rates from app_settings
        logger.info(f"Fetching VAT and commission rates from app_settings")
        settings_res = supabase.table("app_settings").select("value").limit(1).execute()
        settings = {"vat": 0.15, "commission": 0.15}  # defaults
        if settings_res.data and len(settings_res.data) > 0:
            settings_data = settings_res.data[0].get("value")
            if isinstance(settings_data, str):
                settings = json.loads(settings_data)
            else:
                settings = settings_data
        
        vat_percent = settings.get("vat", 0.15) * 100  # Convert to percentage
        commission_percent = settings.get("commission", 0.15)
        logger.info(f"Settings fetched - VAT: {vat_percent}%, Commission: {commission_percent * 100}%")
        
        # Format invoice number: ASK-AP# 04 (based on PO number)
        po_number = po.get("po_number", "")
        invoice_no = f"ASK-AP# {po_number.replace('PO-', '')}"
        logger.info(f"Invoice number formatted: {invoice_no}")
        
        # Create pages
        logger.info(f"Paginating items (30 items per page)")
        pages = InvoiceGeneratorService._paginate_items(items_with_totals)
        logger.info(f"Created {len(pages)} page(s) for invoice")
        
        # Generate multi-page document
        logger.info(f"Generating multi-page invoice document")
        docx_path, pdf_path = InvoiceGeneratorService._generate_multi_page_invoice(
            template_path=template_path,
            output_dir=output_dir,
            pages=pages,
            dealer=dealer,
            po=po,
            invoice_no=invoice_no,
            subtotal=subtotal,
            vat_percent=vat_percent,
            commission_percent=commission_percent,
        )
        
        logger.info(f"Invoice generation completed - DOCX: {docx_path}, PDF: {pdf_path}")
        return docx_path, pdf_path
    
    @staticmethod
    def _paginate_items(items: list) -> list:
        """Split items into pages (30 items per page)."""
        pages = []
        for i in range(0, len(items), ITEMS_PER_PAGE):
            page_items = items[i:i + ITEMS_PER_PAGE]
            pages.append(page_items)
        return pages
    
    @staticmethod
    def _generate_multi_page_invoice(
        template_path: Path,
        output_dir: Path,
        pages: list,
        dealer: dict,
        po: dict,
        invoice_no: str,
        subtotal: float,
        vat_percent: float,
        commission_percent: float,
    ) -> tuple:
        """Generate multi-page invoice document."""
        logger.info(f"Starting multi-page invoice generation with {len(pages)} page(s)")
        logger.debug(f"Template path: {template_path}")
        
        all_docs = []
        balance_bd = 0.0
        total_pages = len(pages)
        
        # Calculate overall VAT and commission (for last page)
        overall_vat = subtotal * (vat_percent / 100)
        overall_commission = subtotal * commission_percent
        logger.info(f"Overall calculations - VAT: {overall_vat}, Commission: {overall_commission}")
        
        for page_num, page_items in enumerate(pages, 1):
            logger.info(f"Processing page {page_num}/{total_pages} with {len(page_items)} items")
            # Add balance b/d as first item if not first page
            if page_num > 1:
                page_items = [
                    {
                        "sl": "B/D",
                        "desc": "Balance B/D",
                        "pkt_size": "",
                        "qty": "",
                        "unit_price": "",
                        "total": balance_bd,
                    }
                ] + page_items
            
            # Calculate page-specific totals (excluding B/D row)
            page_items_only = [item for item in page_items if item.get("sl") != "B/D"]
            page_tp_sum = sum(float(item.get("total", 0)) for item in page_items_only)
            
            # Calculate TIV (Total Including VAT) for this page
            # TIV = sum of TP + (sum of TP * VAT%)
            page_vat = page_tp_sum * (vat_percent / 100)
            page_tiv = page_tp_sum + page_vat
            
            # Add balance B/D to TIV if not first page
            if page_num > 1:
                page_tiv += balance_bd
            
            is_first_page = page_num == 1
            is_last_page = page_num == total_pages
            
            if is_last_page:
                # Last page: show overall totals
                # TIV = overall subtotal (sum of all TP)
                tiv_display = money(subtotal)
                # TEV = TIV - VAT (total excluding vat)
                tev_display = money(subtotal - overall_vat)
                # VAT amount (calculated on overall subtotal)
                vat_display = money(overall_vat)
                # Commission (calculated on overall subtotal)
                commission_display = money(overall_commission)
                # TP = TIV + VAT - Commission (total payable)
                total_payable = subtotal + overall_vat - overall_commission
                tp_display = money(total_payable)
                tp_in_words = number_to_words(total_payable)
                logger.debug(f"Last page - TIV: {subtotal}, VAT: {overall_vat}, Commission: {overall_commission}, TP: {total_payable}")
            else:
                # First/Middle pages: show page TIV (TP + page VAT + balance B/D)
                vat_display = ""
                tev_display = ""
                tiv_display = money(page_tiv)
                commission_display = ""  # Don't show commission on non-last pages
                # TP = page TIV (total payable for this page)
                total_payable = page_tiv
                tp_display = money(total_payable)
                tp_in_words = number_to_words(total_payable)
                logger.debug(f"Page {page_num} - Page TP Sum: {page_tp_sum}, Page VAT: {page_vat}, Balance B/D: {balance_bd}, TIV: {page_tiv}")
            
            # Create context for this page
            context = {
                "COMPANY_NAME": dealer.get("company_name", ""),
                "CONTACT_PERSON": dealer.get("contact_person", ""),
                "CONTACT_NUMBER": dealer.get("contact_number", ""),
                "BILLING_ADDRESS": dealer.get("billing_address", ""),
                "C_CODE": dealer.get("customer_code", ""),
                "INVOICE_NO": invoice_no,
                "DATE": parse_date_ddmmyyyy(po.get("po_date", datetime.now().isoformat())),
                "SHIPPING_ADDRESS": dealer.get("shipping_address", ""),
                "BANK_ACC_NAME": BANK_DETAILS["ac_name"],
                "BANK_ACC_NO": BANK_DETAILS["ac_no"],
                "BANK_BRANCH": BANK_DETAILS["branch_name"],
                "BRANCH_ROUTE_NO": BANK_DETAILS["routing_number"],
                "TEV": tev_display if is_last_page else "",
                "VAT": vat_display,
                "TIV": tiv_display,
                "TP": tp_display,
                "TP_IN_WORDS": tp_in_words,
                "COMM": commission_display,
                "COMMISSION": commission_display,  # Support both placeholder names
                "PAGE_NO": f"{page_num} of {total_pages}",
                "TOTAL_PAGES": str(total_pages),
            }
            
            # Load template and render
            logger.debug(f"Loading template from: {template_path}")
            tpl = DocxTemplate(template_path)
            logger.debug(f"Rendering template with context for page {page_num}")
            logger.info(f"Context COMM value: {context.get('COMM', 'NOT SET')}")
            tpl.render(context)
            buf = BytesIO()
            tpl.save(buf)
            
            doc = Document(BytesIO(buf.getvalue()))
            logger.debug(f"Document created from template for page {page_num}")
            
            # Find and populate items table
            logger.debug(f"Finding items table in document")
            items_table = find_items_table(doc)
            if items_table and len(items_table.rows) >= 2:
                logger.debug(f"Items table found with {len(items_table.rows)} rows")
                template_row = items_table.rows[1]
                
                # Remove existing rows (keep header)
                for row in list(items_table.rows)[1:]:
                    items_table._tbl.remove(row._tr)
                logger.debug(f"Cleared existing rows from items table")
                
                # Add items and fill remaining rows to 30
                logger.debug(f"Adding {len(page_items)} items to table")
                for item in page_items:
                    row_data = [
                        str(item.get("sl", "")),
                        str(item.get("desc", "")),
                        str(item.get("pkt_size", "")),
                        str(item.get("qty", "")) if item.get("qty") != "" else "",
                        money(item.get("unit_price", "")) if item.get("unit_price") != "" else "",
                        money(item.get("total", "")),
                    ]
                    add_table_row_from_template(items_table, row_data, template_row)
                
                # Fill remaining rows with empty rows to make 30 rows total
                rows_to_fill = ITEMS_PER_PAGE - len(page_items)
                if rows_to_fill > 0:
                    logger.debug(f"Filling {rows_to_fill} empty rows to reach {ITEMS_PER_PAGE} rows")
                    for _ in range(rows_to_fill):
                        empty_row_data = ["", "", "", "", "", ""]
                        add_table_row_from_template(items_table, empty_row_data, template_row)
                
                logger.debug(f"All items and empty rows added to table for page {page_num}")
            else:
                logger.warning(f"Items table not found or insufficient rows on page {page_num}")
            
            logger.debug(f"Replacing placeholders in document")
            replace_placeholders_everywhere(doc, context)
            all_docs.append(doc)
            logger.info(f"Page {page_num} document completed and added to collection")
            
            # Calculate balance b/d for next page
            page_total = sum(float(item.get("total", 0)) for item in page_items if item.get("sl") != "B/D")
            balance_bd += page_total
            logger.debug(f"Balance B/D for next page: {balance_bd}")
        
        # Merge all documents
        logger.info(f"Merging {len(all_docs)} document(s)")
        if not all_docs:
            logger.error("No documents were generated")
            raise RuntimeError("No documents generated")
        
        merged_doc = all_docs[0]
        logger.debug(f"Starting with first document as base")
        
        for idx, doc in enumerate(all_docs[1:], 1):
            logger.debug(f"Merging document {idx + 1} into merged document")
            # Append body elements from other documents, excluding section properties
            # Section properties (sectPr) should not be copied to avoid formatting conflicts
            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            sectPr_tag = f'{{{w_ns}}}sectPr'
            
            first_element = True
            for element in doc.element.body:
                # Skip section properties to preserve first page formatting
                if element.tag == sectPr_tag:
                    logger.debug(f"Skipping section properties element to preserve formatting")
                    continue
                
                # For the first element of subsequent documents, add page break to it
                if first_element:
                    element_copy = deepcopy(element)
                    # Add page break to the first element's paragraph properties
                    w_p_tag = f'{{{w_ns}}}p'
                    if element_copy.tag == w_p_tag:
                        pPr = element_copy.find(f'{{{w_ns}}}pPr')
                        if pPr is None:
                            pPr = OxmlElement('w:pPr')
                            element_copy.insert(0, pPr)
                        # Insert page break at the beginning of pPr
                        pageBreak = OxmlElement('w:pageBreakBefore')
                        pPr.insert(0, pageBreak)
                        logger.debug(f"Added page break to first element of page {idx + 1}")
                    first_element = False
                    merged_doc.element.body.append(element_copy)
                else:
                    merged_doc.element.body.append(deepcopy(element))
        
        logger.info(f"All {len(all_docs)} document(s) merged successfully")
        
        # Save DOCX
        safe_no = invoice_no.replace("/", "-").replace("\\", "-").replace("#", "_")
        docx_path = output_dir / f"Invoice_{safe_no}.docx"
        logger.info(f"Saving DOCX file to: {docx_path}")
        merged_doc.save(docx_path)
        logger.info(f"DOCX file saved successfully")
        
        # Convert to PDF
        logger.info(f"Attempting to convert DOCX to PDF")
        pdf_path = convert_docx_to_pdf(docx_path)
        if pdf_path and pdf_path.exists():
            logger.info(f"PDF conversion successful: {pdf_path}")
        else:
            logger.warning(f"PDF conversion failed or LibreOffice not available - will return DOCX instead")
            pdf_path = None
        
        logger.info(f"Invoice generation process completed successfully - DOCX: {docx_path}, PDF: {pdf_path}")
        return docx_path, pdf_path
