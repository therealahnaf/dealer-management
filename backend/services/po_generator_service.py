"""
Purchase Order (PO) generation service.
Generates POs with pagination support.
"""
from pathlib import Path
from io import BytesIO
from datetime import datetime
from copy import deepcopy
import logging

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
from docx.oxml.shared import OxmlElement
from core.database import supabase
from fastapi import HTTPException
from services.utils import convert_docx_to_pdf

logger = logging.getLogger(__name__)

# Create output directory for POs
OUTPUT_DIR = Path(__file__).parent.parent / "output" / "purchase_orders"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
logger.info(f"PO output directory: {OUTPUT_DIR}")

# Hardcoded company details (FROM - ASK INTERNATIONAL)
COMPANY_DETAILS = {
    "company_name": "ASK INTERNATIONAL",
    "contact_person": "Md. Kamrul Hassan",
    "contact_number": "017 6900 4027",
    "billing_address": "F. Haque Tower, 107 CR Dutta Road, Panthapath, Dhaka-1205"
}

# Hardcoded vendor details (TO - ARMY PHARMA)
VENDOR_DETAILS = {
    "company_name": "ARMY PHARMA, BMTF",
    "contact_person": "Mokbul Hossain",
    "contact_number": "01712255785"
}

ITEMS_PER_PAGE = 30


def parse_date_ddmmyyyy(date_str: str) -> str:
    """Parse date string and return in DD/MM/YYYY format."""
    if not date_str:
        return datetime.now().strftime("%d/%m/%Y")
    
    try:
        # Try parsing ISO format (YYYY-MM-DD)
        if 'T' in date_str:
            dt = datetime.fromisoformat(date_str.split('T')[0])
        elif '-' in date_str:
            dt = datetime.strptime(date_str, "%Y-%m-%d")
        elif '/' in date_str:
            # Try MM/DD/YYYY format
            try:
                dt = datetime.strptime(date_str, "%m/%d/%Y")
            except ValueError:
                # Try DD/MM/YYYY format
                dt = datetime.strptime(date_str, "%d/%m/%Y")
        else:
            dt = datetime.now()
        
        return dt.strftime("%d/%m/%Y")
    except Exception as e:
        logger.warning(f"Could not parse date {date_str}: {e}")
        return datetime.now().strftime("%d/%m/%Y")


class POGeneratorService:
    """Generate purchase orders with pagination."""
    
    @staticmethod
    def generate_po_for_dealer(po_id: int, template_path: Path, output_dir: Path = None) -> tuple:
        """
        Generate PO for a purchase order.
        Returns: (docx_path, pdf_path)
        """
        logger.info(f"Starting PO generation for PO ID: {po_id}")
        
        # Use persistent output directory if not specified
        if output_dir is None:
            output_dir = OUTPUT_DIR
        
        output_dir.mkdir(parents=True, exist_ok=True)
        logger.debug(f"Output directory: {output_dir}")
        
        # Get PO details
        logger.info(f"Fetching purchase order details for PO ID: {po_id}")
        po_res = supabase.table("purchase_orders").select("*").eq("po_id", po_id).execute()
        if not po_res.data:
            logger.error(f"Purchase order not found for PO ID: {po_id}")
            raise HTTPException(status_code=404, detail="Purchase order not found")
        
        po = po_res.data[0]
        logger.info(f"PO retrieved: {po.get('po_number')} - Status: {po.get('status')}")
        
        # Get dealer details
        logger.info(f"Fetching dealer details for dealer ID: {po['dealer_id']}")
        dealer_res = supabase.table("dealers").select("*").eq("dealer_id", str(po["dealer_id"])).execute()
        if not dealer_res.data:
            logger.error(f"Dealer not found for dealer ID: {po['dealer_id']}")
            raise HTTPException(status_code=404, detail="Dealer not found")
        
        dealer = dealer_res.data[0]
        logger.info(f"Dealer retrieved: {dealer.get('company_name')}")
        
        # Get PO items
        logger.info(f"Fetching PO items for PO ID: {po_id}")
        items_res = supabase.table("purchase_order_items").select("*").eq("po_id", po_id).execute()
        items = items_res.data if items_res.data else []
        logger.info(f"Retrieved {len(items)} items for PO")
        
        # Fetch product details for each item
        logger.debug("Fetching product details for items")
        for item in items:
            if item.get("product_id"):
                logger.debug(f"Fetching product details for product_id: {item['product_id']}")
                product_res = supabase.table("products").select("name, pack_size").eq("product_id", item["product_id"]).execute()
                if product_res.data:
                    product = product_res.data[0]
                    item["product_name"] = product.get("name", "")  # Use "name" not "product_name"
                    item["pack_size"] = product.get("pack_size", "")
                    logger.debug(f"Product details fetched: name={item['product_name']}, pack_size={item['pack_size']}")
                else:
                    logger.warning(f"No product found for product_id: {item['product_id']}")
                    item["product_name"] = ""
                    item["pack_size"] = ""
            else:
                logger.debug(f"Item has no product_id: {item}")
                item["product_name"] = ""
                item["pack_size"] = ""
        
        # Paginate items
        pages = []
        for i in range(0, len(items), ITEMS_PER_PAGE):
            pages.append(items[i:i + ITEMS_PER_PAGE])
        
        if not pages:
            pages = [[]]
        
        logger.info(f"PO will have {len(pages)} page(s)")
        
        # Generate multi-page PO
        docx_path, pdf_path = POGeneratorService._generate_multi_page_po(
            po=po,
            dealer=dealer,
            pages=pages,
            template_path=template_path,
            output_dir=output_dir
        )
        
        logger.info(f"PO generation completed - DOCX: {docx_path}, PDF: {pdf_path}")
        return docx_path, pdf_path
    
    @staticmethod
    def _generate_multi_page_po(po: dict, dealer: dict, pages: list, template_path: Path, output_dir: Path) -> tuple:
        """Generate multi-page PO document."""
        logger.info(f"Generating multi-page PO with {len(pages)} page(s)")
        
        total_pages = len(pages)
        po_number = po.get("po_number", "")
        po_ref = dealer.get("company_name", "")
        
        all_docs = []
        
        for page_num, page_items in enumerate(pages, 1):
            logger.info(f"Processing page {page_num} of {total_pages}")
            
            # Create context for this page
            context = {
                # TO (Vendor)
                "VENDOR_COMPANY_NAME": VENDOR_DETAILS["company_name"],
                "VENDOR_CONTACT_PERSON": VENDOR_DETAILS["contact_person"],
                "VENDOR_CONTACT_NUMBER": VENDOR_DETAILS["contact_number"],
                # PO Details
                "PO_NO": po_number,
                "DATE": parse_date_ddmmyyyy(po.get("po_date", datetime.now().isoformat())),
                "PAGE_NO": f"{page_num} of {total_pages}",
                # PO Reference
                "PO_REF": po_ref,
                # FROM (Company) - These match template placeholders exactly
                "COMPANY_NAME": COMPANY_DETAILS["company_name"],
                "CONTACT_PERSON": COMPANY_DETAILS["contact_person"],
                "CONTACT_NUMBER": COMPANY_DETAILS["contact_number"],
                "BILLING_ADDR": COMPANY_DETAILS["billing_address"],
                "SHIPPING_ADDR": COMPANY_DETAILS["billing_address"],
                # Dealer Shipping Address
                "DEALER_SHIPPING_ADDR": dealer.get("shipping_address", ""),
            }
            
            logger.debug(f"Context created: {context}")
            
            # Load template and render
            logger.debug(f"Loading template from: {template_path}")
            tpl = DocxTemplate(template_path)
            logger.debug(f"Rendering template with context for page {page_num}")
            logger.debug(f"Context keys: {list(context.keys())}")
            logger.debug(f"Context values: {context}")
            
            try:
                tpl.render(context)
                logger.debug(f"Template rendered successfully")
            except Exception as e:
                logger.error(f"Error rendering template: {e}", exc_info=True)
            
            buf = BytesIO()
            tpl.save(buf)
            
            doc = Document(BytesIO(buf.getvalue()))
            logger.debug(f"Document created from template for page {page_num}")
            
            # Find and populate items table
            logger.debug(f"Finding items table in document")
            items_table = POGeneratorService._find_items_table(doc)
            if items_table and len(items_table.rows) >= 2:
                logger.debug(f"Items table found with {len(items_table.rows)} rows")
                template_row = items_table.rows[1]
                
                # Remove existing rows (keep header)
                for row in list(items_table.rows)[1:]:
                    items_table._tbl.remove(row._tr)
                logger.debug(f"Cleared existing rows from items table")
                
                # Add items and fill remaining rows to 30
                logger.debug(f"Adding {len(page_items)} items to table")
                for idx, item in enumerate(page_items, 1):
                    # Get product name from item (should be populated from product details fetch)
                    product_name = item.get("product_name", "")
                    pack_size = item.get("pack_size", "")
                    quantity = item.get("quantity", "")
                    
                    logger.debug(f"Item {idx}: product_name='{product_name}', pack_size='{pack_size}', quantity='{quantity}'")
                    
                    row_data = [
                        str(idx),  # Sl #
                        po_number,  # Invoice # (use PO_NO)
                        parse_date_ddmmyyyy(po.get("po_date", "")),  # PO Date
                        product_name,  # Product Name
                        str(pack_size),  # Pkt Size
                        str(quantity),  # Qty
                    ]
                    logger.debug(f"Row data: {row_data}")
                    POGeneratorService._add_table_row_from_template(items_table, row_data, template_row)
                
                # Fill remaining rows with empty rows to make 30 rows total
                rows_to_fill = ITEMS_PER_PAGE - len(page_items)
                if rows_to_fill > 0:
                    logger.debug(f"Filling {rows_to_fill} empty rows to reach {ITEMS_PER_PAGE} rows")
                    for _ in range(rows_to_fill):
                        empty_row_data = ["", "", "", "", "", ""]
                        POGeneratorService._add_table_row_from_template(items_table, empty_row_data, template_row)
                
                logger.debug(f"All items and empty rows added to table for page {page_num}")
            else:
                logger.warning(f"Items table not found or insufficient rows on page {page_num}")
            
            # Replace any remaining placeholders in the document
            logger.debug(f"Replacing placeholders in document")
            result = POGeneratorService._replace_placeholders_everywhere(doc, context)
            logger.debug(f"Placeholder replacement result: {result}")
            
            all_docs.append(doc)
            logger.info(f"Page {page_num} document completed and added to collection")
        
        # Merge all documents
        logger.info(f"Merging {len(all_docs)} document(s)")
        if not all_docs:
            logger.error("No documents were generated")
            raise RuntimeError("No documents generated")
        
        merged_doc = all_docs[0]
        logger.debug(f"Starting with first document as base")
        
        for idx, doc in enumerate(all_docs[1:], 1):
            logger.debug(f"Merging document {idx + 1} into merged document")
            # Append body elements from other documents, excluding section properties
            # Section properties (sectPr) should not be copied to avoid formatting conflicts
            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            sectPr_tag = f'{{{w_ns}}}sectPr'
            
            first_element = True
            for element in doc.element.body:
                # Skip section properties to preserve first page formatting
                if element.tag == sectPr_tag:
                    logger.debug(f"Skipping section properties element to preserve formatting")
                    continue
                
                # For the first element of subsequent documents, add page break to it
                if first_element:
                    element_copy = deepcopy(element)
                    # Add page break to the first element's paragraph properties
                    w_p_tag = f'{{{w_ns}}}p'
                    if element_copy.tag == w_p_tag:
                        pPr = element_copy.find(f'{{{w_ns}}}pPr')
                        if pPr is None:
                            pPr = OxmlElement('w:pPr')
                            element_copy.insert(0, pPr)
                        # Insert page break at the beginning of pPr
                        pageBreak = OxmlElement('w:pageBreakBefore')
                        pPr.insert(0, pageBreak)
                        logger.debug(f"Added page break to first element of page {idx + 1}")
                    first_element = False
                    merged_doc.element.body.append(element_copy)
                else:
                    merged_doc.element.body.append(deepcopy(element))
        
        # Save DOCX
        docx_filename = f"PO_{po_number}.docx"
        docx_path = output_dir / docx_filename
        logger.info(f"Saving DOCX to: {docx_path}")
        merged_doc.save(docx_path)
        logger.info(f"DOCX saved successfully")
        
        # Try to convert to PDF
        pdf_path = convert_docx_to_pdf(docx_path)
        
        return docx_path, pdf_path
    
    @staticmethod
    def _replace_placeholders_everywhere(doc: Document, context: dict):
        """Replace placeholders in all text nodes."""
        try:
            repl = {f"{{{{ {k} }}}}": str(v) for k, v in context.items()}
            repl.update({f"{{{{{k}}}}}": str(v) for k, v in context.items()})

            W_T = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
            root = doc.part.element

            for t in root.iter(W_T):
                if t.text:
                    orig = t.text
                    new = orig
                    for k, v in repl.items():
                        if k in new:
                            new = new.replace(k, v)
                    if new != orig:
                        t.text = new
            return True
        except Exception as e:
            logger.error(f"Error replacing placeholders: {e}")
            return False
    
    @staticmethod
    def _find_items_table(doc: Document):
        """Find the items table by matching header cells."""
        wanted = ["Sl #", "Invoice #", "PO Date", "Product Name", "Pkt Size", "Qty"]
        for tbl in doc.tables:
            try:
                hdr = [cell.text.strip() for cell in tbl.rows[0].cells]
                if len(hdr) >= 6 and all(h in hdr for h in wanted):
                    return tbl
            except Exception:
                pass
        # Fallback: return second table if headers don't match
        tables = doc.tables
        if len(tables) >= 2:
            return tables[1]
        elif len(tables) == 1:
            return tables[0]
        return None
    
    @staticmethod
    def _add_table_row_from_template(table, row_data: list, template_row):
        """Clone template row and replace text, preserving formatting (exact invoice implementation)."""
        def set_cell_font(cell, font_name, font_size):
            """Set font for all runs in a cell."""
            try:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = Pt(font_size)
            except Exception:
                pass
        
        try:
            w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            W_T = f'{{{w_ns}}}t'
            W_P = f'{{{w_ns}}}p'
            W_R = f'{{{w_ns}}}r'
            W_TC = f'{{{w_ns}}}tc'

            new_tr = deepcopy(template_row._element)
            table._tbl.append(new_tr)

            cells = new_tr.findall(f'.//{W_TC}')
            for idx, cell_text in enumerate(row_data):
                if idx >= len(cells):
                    break
                cell = cells[idx]
                t_elems = list(cell.iter(W_T))

                if t_elems:
                    t_elems[0].text = str(cell_text)
                    for t in t_elems[1:]:
                        t.text = ''
                else:
                    first_p = cell.find(f'.//{W_P}')
                    if first_p is None:
                        first_p = OxmlElement('w:p')
                        cell.append(first_p)
                    first_r = first_p.find(f'.//{W_R}')
                    if first_r is None:
                        first_r = OxmlElement('w:r')
                        first_p.append(first_r)
                    t = OxmlElement('w:t')
                    t.text = str(cell_text)
                    first_r.append(t)

            new_row = table.rows[-1]
            for cell in new_row.cells:
                set_cell_font(cell, 'Arial', 8)

            return True
        except Exception as e:
            logger.error(f"Error adding row: {e}")
            return False
    
