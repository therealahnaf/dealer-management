from pathlib import Path
from io import BytesIO
from datetime import datetime
from copy import deepcopy
import re
import subprocess

from docxtpl import DocxTemplate
from docx import Document
from docx.oxml.shared import OxmlElement
from docx.shared import Pt


# -----------------------------
# HELPERS
# -----------------------------
def parse_date_mmddyyyy(s: str) -> str:
    """Normalize date to MM/DD/YYYY if possible, cross-platform safe."""
    for fmt in ("%m/%d/%Y", "%m/%d/%y", "%Y-%m-%d"):
        try:
            dt = datetime.strptime(s, fmt)
            return dt.strftime("%m/%d/%Y")
        except Exception:
            continue
    try:
        parts = s.replace("-", "/").split("/")
        if len(parts) == 3:
            m, d, y = parts
            dt = datetime.strptime(f"{int(m)}/{int(d)}/{int(y)}", "%m/%d/%Y")
            return dt.strftime("%m/%d/%Y")
    except Exception:
        pass
    return s


def money(x) -> str:
    """Format number with thousand separators."""
    try:
        f = float(x)
        return f"{int(f):,}" if f.is_integer() else f"{f:,.2f}"
    except Exception:
        return str(x)


def number_to_words(num):
    """Convert number to words (basic implementation for Bangladeshi Taka)."""
    try:
        num = float(num)
        taka = int(num)
        paisa = int(round((num - taka) * 100))

        ones = ['', 'One', 'Two', 'Three', 'Four', 'Five', 'Six', 'Seven', 'Eight', 'Nine']
        teens = ['Ten', 'Eleven', 'Twelve', 'Thirteen', 'Fourteen', 'Fifteen',
                 'Sixteen', 'Seventeen', 'Eighteen', 'Nineteen']
        tens = ['', '', 'Twenty', 'Thirty', 'Forty', 'Fifty', 'Sixty', 'Seventy', 'Eighty', 'Ninety']

        def convert_below_thousand(n):
            if n == 0:
                return ''
            if n < 10:
                return ones[n]
            if n < 20:
                return teens[n - 10]
            if n < 100:
                return tens[n // 10] + (' ' + ones[n % 10] if n % 10 else '')
            return ones[n // 100] + ' Hundred' + (' ' + convert_below_thousand(n % 100) if n % 100 else '')

        if taka == 0:
            result = 'Zero Taka'
        else:
            crore = taka // 10000000
            lakh = (taka % 10000000) // 100000
            thousand = (taka % 100000) // 1000
            hundred = taka % 1000

            parts = []
            if crore:   parts.append(convert_below_thousand(crore) + ' Crore')
            if lakh:    parts.append(convert_below_thousand(lakh) + ' Lakh')
            if thousand: parts.append(convert_below_thousand(thousand) + ' Thousand')
            if hundred: parts.append(convert_below_thousand(hundred))
            result = (' '.join(parts) + ' Taka').strip()

        if paisa:
            result += ' and ' + convert_below_thousand(paisa) + ' Paisa'

        return result.strip() + ' Only'
    except Exception:
        return "Amount in Words"


def set_cell_font(cell, font_name='Arial', font_size=8):
    """Set font name and size for all runs in a cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)


def add_table_row_from_template(table, row_data: list, template_row):
    """
    Clone the template row and replace only the text in existing runs,
    preserving all formatting (font size, bold, alignment, borders, etc.).
    Then apply Arial 8pt font to all cells.
    """
    try:
        w_ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        W_T = f'{{{w_ns}}}t'
        W_P = f'{{{w_ns}}}p'
        W_R = f'{{{w_ns}}}r'
        W_TC = f'{{{w_ns}}}tc'

        # 1) Deep copy the template row (with formatting) and append it
        new_tr = deepcopy(template_row._element)
        table._tbl.append(new_tr)

        # 2) For each cell, overwrite only the first <w:t> text; blank out the rest
        cells = new_tr.findall(f'.//{W_TC}')
        for idx, cell_text in enumerate(row_data):
            if idx >= len(cells):
                break
            cell = cells[idx]

            t_elems = list(cell.iter(W_T))

            if t_elems:
                t_elems[0].text = str(cell_text)
                for t in t_elems[1:]:
                    t.text = ''
            else:
                first_p = cell.find(f'.//{W_P}')
                if first_p is None:
                    first_p = OxmlElement('w:p')
                    cell.append(first_p)
                first_r = first_p.find(f'.//{W_R}')
                if first_r is None:
                    first_r = OxmlElement('w:r')
                    first_p.append(first_r)
                t = OxmlElement('w:t')
                t.text = str(cell_text)
                first_r.append(t)

        # 3) Now apply Arial 8pt to all cells in the newly added row
        new_row = table.rows[-1]
        for cell in new_row.cells:
            set_cell_font(cell, 'Arial', 8)

        return True
    except Exception as e:
        print(f"Error adding row (format-preserving): {e}")
        import traceback; traceback.print_exc()
        return False


def find_items_table(doc: Document):
    """Locate the items table by matching header cells."""
    wanted = ["Sl #", "Product Description", "Pkt Size", "Qty", "Unit Price", "Total Price"]
    for tbl in doc.tables:
        try:
            hdr = [cell.text.strip() for cell in tbl.rows[0].cells]
            if len(hdr) >= 6 and all(h in hdr for h in wanted):
                return tbl
        except Exception:
            pass
    return None


def find_and_fill_summary_table(doc: Document, context: dict, commission_rate: float):
    """Fill the financial summary table by label matching (works for nested tables)."""
    try:
        tev = float(str(context['TEV']).replace(',', ''))
        commission = tev * commission_rate
        advance = float(str(context.get('ADVANCE', 0)).replace(',', ''))
        previous_due = float(str(context.get('PREVIOUS_DUE', 0)).replace(',', ''))

        def process_table(table):
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    text = cell.text.strip().lower()
                    for nested in cell.tables:
                        process_table(nested)

                    if 'commission @' in text or 'commission@' in text:
                        j = i + 1 if i + 1 < len(row.cells) else -1
                        row.cells[j].text = money(commission)
                    elif 'advance' in text and ('ii)' in text or 'advance:' in text):
                        j = i + 1 if i + 1 < len(row.cells) else -1
                        row.cells[j].text = money(advance)
                    elif 'previous due' in text and ('i)' in text or 'previous due:' in text):
                        j = i + 1 if i + 1 < len(row.cells) else -1
                        row.cells[j].text = money(previous_due)

        for t in doc.tables:
            process_table(t)
        return True
    except Exception as e:
        print(f"Error filling summary table: {e}")
        import traceback; traceback.print_exc()
        return False


def replace_placeholders_everywhere(doc: Document, context: dict):
    """
    Sweep remaining placeholders in all text nodes (including shapes) WITHOUT xpath namespaces.
    Note: won't fix placeholders that are split across runs.
    """
    try:
        repl = {f"{{{{ {k} }}}}": str(v) for k, v in context.items()}
        repl.update({f"{{{{{k}}}}}": str(v) for k, v in context.items()})

        W_T = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
        root = doc.part.element

        changed = 0
        for t in root.iter(W_T):
            if t.text:
                orig = t.text
                new = orig
                for k, v in repl.items():
                    if k in new:
                        new = new.replace(k, v)
                if new != orig:
                    t.text = new
                    changed += 1
        print(f"Sweep replaced text nodes: {changed}")
        return True
    except Exception as e:
        print(f"Error sweeping placeholders (iter): {e}")
        import traceback; traceback.print_exc()
        return False


def list_unresolved_placeholders(doc: Document):
    """
    Scan for any {{...}} tokens still present (diagnostic).
    """
    W_T = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
    token_re = re.compile(r'{{[^}]+}}')
    found = set()
    for t in doc.part.element.iter(W_T):
        if t.text:
            for m in token_re.findall(t.text):
                found.add(m)
    if found:
        print("Unresolved placeholders still in document:", sorted(found))
    else:
        print("No unresolved placeholders detected.")


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path | None = None) -> Path | None:
    """Convert DOCX to PDF using LibreOffice (if available)."""
    if pdf_path is None:
        pdf_path = docx_path.with_suffix('.pdf')
    try:
        subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf',
             '--outdir', str(pdf_path.parent), str(docx_path)],
            check=True, capture_output=True, text=True
        )
        print("PDF converted using LibreOffice")
        return pdf_path
    except (subprocess.CalledProcessError, FileNotFoundError):
        print("LibreOffice not found. Install it for PDF conversion.")
        return None


class InvoiceService:
    def __init__(self, template_path: Path, output_dir: Path):
        self.template_path = template_path
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate_invoice(self, data: dict, generate_pdf: bool = True):
        # Calculate totals
        items_with_totals = []
        subtotal_tev = 0.0

        for item in data["items"]:
            line_total = float(item["qty"]) * float(item["unit_price"])
            subtotal_tev += line_total
            items_with_totals.append({
                "sl": item.get("sl"),
                "desc": item.get("desc", ""),
                "pkt_size": item.get("pkt_size", ""),
                "qty": item.get("qty", 0),
                "unit_price": item.get("unit_price", 0.0),
                "total": line_total,
            })

        vat_rate = float(data.get("vat_percent", 0)) / 100.0
        commission_rate = float(data.get("commission_percent", 0)) / 100.0

        vat_amount = subtotal_tev * vat_rate
        tiv = subtotal_tev + vat_amount
        commission = subtotal_tev * commission_rate
        advance = float(data.get("advance", 0) or 0)
        previous_due = float(data.get("previous_due", 0) or 0)
        total_payable = tiv - commission - advance + previous_due

        context = {
            "COMPANY_NAME": data["customer"]["dealer_name"],
            "CONTACT_PERSON": data["customer"]["customer_name"],
            "CONTACT_NUMBER": data["customer"]["contact_number"],
            "BILLING_ADDRESS": data["customer"]["billing_address"],
            "C_CODE": data["customer"]["code"],
            "INVOICE_NO": data["invoice_no"],
            "DATE": parse_date_mmddyyyy(data["date"]),
            "SHIPPING_ADDRESS": data["customer"]["shipping_address"],
            "BANK_ACC_NAME": data["bank"]["ac_name"],
            "BANK_ACC_NO": data["bank"]["ac_no"],
            "BANK_BRANCH": data["bank"]["branch_name"],
            "BRANCH_ROUTE_NO": data["bank"]["routing_number"],
            "TEV": money(subtotal_tev),
            "VAT": money(vat_amount),
            "TIV": money(tiv),
            "COMMISSION": money(commission),
            "ADVANCE": money(advance),
            "PREVIOUS_DUE": money(previous_due),
            "TP": money(total_payable),
            "TP_IN_WORDS": number_to_words(total_payable),
            "PAGE_NO": "1",
        }

        tpl = DocxTemplate(self.template_path)
        tpl.render(context)
        buf = BytesIO()
        tpl.save(buf)

        doc = Document(BytesIO(buf.getvalue()))

        items_table = find_items_table(doc)
        if items_table is None:
            raise RuntimeError("Items table not found in template.")
        if len(items_table.rows) < 2:
            raise RuntimeError("Template must have at least 2 rows (header + template row) in items table.")
        template_row = items_table.rows[1]

        for row in list(items_table.rows)[1:]:
            items_table._tbl.remove(row._tr)

        for item in items_with_totals:
            row_data = [
                str(item.get("sl", "")),
                str(item.get("desc", "")),
                str(item.get("pkt_size", "")),
                str(item.get("qty", "")),
                money(item.get("unit_price", "")) if item.get("unit_price") != "" else "",
                money(item.get("total", "")),
            ]
            if not add_table_row_from_template(items_table, row_data, template_row):
                print(f"Failed to add row {item.get('sl')}")

        find_and_fill_summary_table(doc, context, commission_rate)
        replace_placeholders_everywhere(doc, context)
        list_unresolved_placeholders(doc)

        safe_no = data["invoice_no"].replace("/", "-").replace("\\", "-").replace("#", "_")
        docx_path = self.output_dir / f"Invoice_{safe_no}.docx"
        doc.save(docx_path)

        pdf_path = None
        if generate_pdf:
            pdf_path = convert_docx_to_pdf(docx_path)

        return docx_path, pdf_path
